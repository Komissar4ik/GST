import re
import sys
from typing import Optional

SPECIAL_HEADINGS: dict[str, str] = {
    'введение':                         'ВВЕДЕНИЕ',
    'заключение':                       'ЗАКЛЮЧЕНИЕ',
    'содержание':                       'СОДЕРЖАНИЕ',
    'оглавление':                       'ОГЛАВЛЕНИЕ',
    'список литературы':                'СПИСОК ЛИТЕРАТУРЫ',
    'список использованных источников': 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ',
    'библиографический список':         'БИБЛИОГРАФИЧЕСКИЙ СПИСОК',
    'список сокращений':                'СПИСОК СОКРАЩЕНИЙ',
    'список обозначений':               'СПИСОК ОБОЗНАЧЕНИЙ',
    'список терминов':                  'СПИСОК ТЕРМИНОВ',
    'аннотация':                        'АННОТАЦИЯ',
    'реферат':                          'РЕФЕРАТ',
    'abstract':                         'ABSTRACT',
    'благодарности':                    'БЛАГОДАРНОСТИ',
    'глоссарий':                        'ГЛОССАРИЙ',
}

_APPENDIX_RE = re.compile(r'^приложение\s*(.*)$', re.IGNORECASE)
_EXISTING_NUM_RE = re.compile(r'^(\d+\.)*\d+\.?\s*')
_TRAILING_PUNCT_RE = re.compile(r'[.,;:]\s*$')
_HEADING_RE = re.compile(r'^(#{1,6})\s+(.*)')


def _normalize(raw: str) -> str:
    t = _EXISTING_NUM_RE.sub('', raw.strip())
    return _TRAILING_PUNCT_RE.sub('', t).strip()


def _key(title: str) -> str:
    return _normalize(title).lower()


def is_special(title: str) -> bool:
    k = _key(title)
    if k in SPECIAL_HEADINGS:
        return True
    if _APPENDIX_RE.match(k):
        return True
    return False


def canonical_form(title: str) -> str:
    k = _key(title)
    if k in SPECIAL_HEADINGS:
        return SPECIAL_HEADINGS[k]
    m = _APPENDIX_RE.match(k)
    if m:
        suffix = m.group(1).strip().upper()
        return f'ПРИЛОЖЕНИЕ {suffix}' if suffix else 'ПРИЛОЖЕНИЕ'
    return _normalize(title)


def fix_special_headings_text(text: str, force_uppercase: bool = True) -> str:
    lines_out = []
    for line in text.splitlines():
        m = _HEADING_RE.match(line)
        if m:
            hashes, title = m.group(1), m.group(2)
            if is_special(title):
                fixed = canonical_form(title) if force_uppercase else _normalize(title)
                lines_out.append(f'{hashes} {fixed}')
            else:
                lines_out.append(line)
        else:
            lines_out.append(line)
    return '\n'.join(lines_out)


def _set_para_text(para, new_text: str) -> None:
    for run in para.runs[1:]:
        run._element.getparent().remove(run._element)
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def fix_special_headings_docx(input_path: str,
                               output_path: Optional[str] = None,
                               force_uppercase: bool = True) -> None:
    try:
        from docx import Document
    except ImportError:
        print('pip install python-docx', file=sys.stderr)
        sys.exit(1)

    doc = Document(input_path)
    changed = 0
    for para in doc.paragraphs:
        if not para.style.name.startswith('Heading'):
            continue
        raw = para.text.strip()
        if not is_special(raw):
            continue
        new_text = canonical_form(raw) if force_uppercase else _normalize(raw)
        if new_text != raw:
            _set_para_text(para, new_text)
            changed += 1

    save_path = output_path or input_path
    doc.save(save_path)
    print(f'Нормализовано: {changed}. Сохранено: {save_path}')
