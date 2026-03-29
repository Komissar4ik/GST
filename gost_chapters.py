import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import sys

def get_chapter_type(paragraph):
    """
    Определяет тип главы:
    - "gost": правильная глава ("ГЛАВА 1. ...")
    - "bad": глава без слова "ГЛАВА" (например, "1 ТЕКСТ")
    - "no_dot": глава без точки ("ГЛАВА 1 Анализ ...")
    - None: не глава
    """
    text = paragraph.text.strip()
    uptext = text.upper()
    if re.match(r'^ГЛАВА\s+\d+\.\s+.+', uptext):
        return "gost"
    elif re.match(r'^\d+\s+[A-ZА-ЯЁ]', uptext) and not re.match(r'^\d+\.\d+', uptext):
        return "bad"
    elif re.match(r'^ГЛАВА\s+\d+\s+[^\d\W]', text, re.IGNORECASE):
        return "no_dot"
    else:
        return None

def fix_chapter(paragraph, num, chapter_type):
    """
    Приводит разные типы глав к формату ГОСТ.
    """
    text = paragraph.text.strip()
    if chapter_type == "bad":
        # Было "1 текст" → ГЛАВА <num>. Текст
        m = re.match(r'^(\d+)\s+(.+)$', text)
        if m:
            title = m.group(2).capitalize()
            new_text = f'ГЛАВА {num}. {title}'
            apply_heading_style(paragraph, new_text)
    elif chapter_type == "no_dot":
        # Было "ГЛАВА 1 Анализ ..." → ГЛАВА 1. Анализ
        m = re.match(r'^(ГЛАВА)\s+(\d+)\s+(.+)$', text, re.IGNORECASE)
        if m:
            title = m.group(3).capitalize()
            new_text = f'ГЛАВА {num}. {title}'
            apply_heading_style(paragraph, new_text)

def apply_heading_style(paragraph, text):
    """
    Применяет ГОСТ-стиль к абзацу (жирный, Times New Roman 14, без абзацного отступа, выравнивание влево).
    """
    try:
        paragraph.style = 'Normal'
    except:
        paragraph.style = 'Обычный'
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def is_section(paragraph):
    # Проверяет: это подзаголовок раздела ("1.1 ...", "1.2.1 ...", и т.д.)
    return re.match(r'^(\d+(\.\d+)+)\s+\S', paragraph.text.strip()) is not None

def get_section_number(paragraph):
    # Возвращает номер раздела в виде строки и списка целых чисел
    m = re.match(r'^((\d+(?:\.\d+)+))\s+\S', paragraph.text.strip())
    num_str = m.group(1) if m else None
    num_list = list(map(int, num_str.split('.'))) if num_str else None
    return num_str, num_list

def strict_sequence_check(number_lists, current):
    # Проверяет строгую последовательность подразделов (например, 1.1, 1.2, 1.3 ...)
    same_level = [nums for nums in number_lists if len(nums) == len(current) and nums[:-1] == current[:-1]]
    if not same_level:
        return current[-1] == 1
    prev = max(same_level)
    return current[-1] == prev[-1] + 1

def format_document_headings_strict(docx_path, output_path):
    """
    Открывает входной docx, проходит по каждому абзацу, и приводит главы/подглавы к ГОСТ-формату.
    """
    doc = Document(docx_path)

    in_main_part = False
    prev_was_chapter = False
    current_chapter = None
    all_section_numbers = []
    expected_chapter_num = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        uptext = text.upper()

        if uptext == "ВВЕДЕНИЕ":
            in_main_part = True
            current_chapter = None
            prev_was_chapter = False
            all_section_numbers.clear()
            expected_chapter_num = 1
            continue
        if uptext == "ЗАКЛЮЧЕНИЕ":
            in_main_part = False

        if in_main_part:
            chapter_type = get_chapter_type(para)

            if chapter_type == "bad" or chapter_type == "no_dot":
                fix_chapter(para, expected_chapter_num, chapter_type)
                current_chapter = expected_chapter_num
                prev_was_chapter = True
                all_section_numbers.clear()
                expected_chapter_num += 1
                continue

            if chapter_type == "gost":
                # Проверить, совпадает ли номер главы с ожидаемым
                m = re.match(r'^ГЛАВА\s+(\d+)', uptext)
                if m:
                    corr_num = int(m.group(1))
                    title = re.sub(r'^ГЛАВА\s+\d+\.?', '', para.text, flags=re.IGNORECASE).strip()
                    if corr_num != expected_chapter_num:
                        new_text = f'ГЛАВА {expected_chapter_num}. {title.capitalize()}'
                        apply_heading_style(para, new_text)
                    else:
                        new_text = f'ГЛАВА {expected_chapter_num}. {title.capitalize()}'
                        apply_heading_style(para, new_text)
                    current_chapter = expected_chapter_num
                else:
                    current_chapter = expected_chapter_num
                prev_was_chapter = True
                all_section_numbers.clear()
                expected_chapter_num += 1
                continue

            # 4. Проверяем подзаголовки разделов
            if is_section(para):
                num_str, num_list = get_section_number(para)
                if not num_list:
                    continue
                if prev_was_chapter:
                    if num_list[-1] == 1:
                        apply_heading_style(para, para.text.strip())
                        all_section_numbers.append(num_list)
                    prev_was_chapter = False
                    continue
                if strict_sequence_check(all_section_numbers, num_list):
                    apply_heading_style(para, para.text.strip())
                    all_section_numbers.append(num_list)
                continue

            prev_was_chapter = False

    doc.save(output_path)

if __name__ == "__main__":
    # Универсальный запуск с аргументами командной строки
    if len(sys.argv) < 3:
        print("Использование: python gost_chapters.py <входной_файл.docx> <выходной_файл.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]
    format_document_headings_strict(input_file, output_file)
    print(f"Готово! Результат сохранён в {output_file}")